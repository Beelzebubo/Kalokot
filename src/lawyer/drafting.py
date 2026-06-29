"""Complaint / RTI / FOIA draft generator.

Generates ready-to-file legal documents (complaint letters, RTI requests,
whistleblower reports) either via LLM (tailored) or from YAML templates
(fallback).  Also provides plain-text and .docx export utilities.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from ..shared.models import JurisdictionCode, ComplaintDraft, CounselRequest
from ..shared.jurisdiction import JurisdictionLoader
from ..shared.llm import LLMClient


# ── LLM system prompt for legal-document drafting ─────────────────────────

DRAFTING_SYSTEM_PROMPT = """You are a legal document drafting assistant specializing in public procurement complaints, 
whistleblower reports, and RTI/FOIA requests. You have access to a legal knowledge base 
with jurisdiction-specific templates.

Your job:
1. Given a user's description of a procurement violation and the relevant jurisdiction,
2. Generate a complete, ready-to-file legal document.
3. Follow the structure and language conventions of that jurisdiction.

Rules:
- Use the jurisdiction's templates as a base structure
- Fill in specific details from the user's complaint
- Keep the language formal and professional
- Include placeholders [in brackets] for user-specific fields
- Add practical filing instructions at the bottom
- Clearly note: "This is a DRAFT. Review with a qualified attorney before filing."

Output format: plain text document with clear section headings."""


class DraftGenerator:
    """Generate complaint / RTI / whistleblower draft documents.

    Two modes:
    * **LLM mode** — uses the configured LLM client to produce a tailored
      document based on the user's tender context and question.
    * **Template mode** — selects a YAML template by keyword matching on the
      question text, fills in basic structure.

    Also exposes :meth:`export_txt` and :meth:`export_docx` for output.
    """

    def __init__(self, loader: JurisdictionLoader, llm: Optional[LLMClient] = None):
        """Initialise the draft generator.

        Args:
            loader: Jurisdiction loader (used for legal context + templates).
            llm: Optional LLM client for AI-generated drafts.
        """
        self.loader = loader
        self.llm = llm

    def generate_complaint(self, request: CounselRequest) -> ComplaintDraft:
        """Generate a complaint draft based on the counsel request.

        Delegates to LLM if available, otherwise falls back to template
        selection.

        Args:
            request: The counsel request (tender context, question, risk report).

        Returns:
            A ComplaintDraft with title, body, template_name, and instructions.
        """
        if self.llm:
            return self._generate_with_llm(request)
        else:
            return self._generate_from_template(request)

    def _generate_with_llm(self, request: CounselRequest) -> ComplaintDraft:
        """Use LLM to generate a tailored complaint draft.

        Builds a prompt containing the jurisdiction, tender description, user
        question, available template names, and relevant legal context.
        """
        jurisdiction = request.jurisdiction
        legal_context = self.loader.build_legal_context(
            jurisdiction,
            flagged_ids=[f.red_flag_id for f in (request.risk_report.flagged_clauses
                                                 if request.risk_report else [])]
            if request.risk_report else None
        )

        # Collect available template names for the LLM to reference
        templates = self.loader.get_templates(jurisdiction)
        template_names = [t.get("template_name") or t.get("name") or t.get("id", "?")
                          for t in templates]

        user_prompt = (
            f"I need a legal document regarding a procurement violation.\n\n"
            f"Jurisdiction: {jurisdiction.value}\n\n"
            f"Tender Description: {request.tender_context[:2000]}\n\n"
            f"User's Specific Question/Request: {request.question}\n\n"
            f"Available Templates: {', '.join(template_names)}\n\n"
            f"Legal Context:\n{legal_context[:3000]}\n\n"
            f"Please generate the most appropriate legal document based on the user's request. "
            f"If the user didn't specify a document type, generate a complaint letter to the "
            f"appropriate oversight body."
        )

        draft_text = self.llm.generate(
            system_prompt=DRAFTING_SYSTEM_PROMPT,
            user_prompt=user_prompt,
            temperature=0.3,
            max_tokens=4096,
        )

        # Build a human-readable title from the first 50 chars of tender context
        if len(request.tender_context) > 50:
            title = f"Complaint Draft — {request.tender_context[:50]}..."
        else:
            title = f"Complaint Draft — {request.tender_context}"

        return ComplaintDraft(
            title=title,
            jurisdiction=jurisdiction,
            body=draft_text,
            template_name="llm_generated",
            instructions="Review carefully with a qualified attorney before filing.",
        )

    def _generate_from_template(self, request: CounselRequest) -> ComplaintDraft:
        """Fallback: generate a draft from a pre-defined YAML template.

        Selects a template based on keyword matching against the user's
        question text (e.g. "ciaa" → complaint_ciaa, "rti" → rti_request).
        """
        jurisdiction = request.jurisdiction
        question_lower = request.question.lower()

        # ── Detect intent from question keywords ───────────────────────────
        if "kpk" in question_lower or "ciaa" in question_lower or "anti-corruption" in question_lower:
            template_name = "complaint_ciaa"
        elif "rti" in question_lower or "right to information" in question_lower or "foia" in question_lower or "information" in question_lower:
            template_name = "rti_request"
        elif "appeal" in question_lower or "disqualif" in question_lower:
            template_name = "appeal_review_committee"
        elif "ppmo" in question_lower:
            template_name = "complaint_ppmo"
        else:
            template_name = "complaint_ciaa"

        # Load the selected template body from the jurisdiction YAML files
        template = self.loader.get_template_by_name(jurisdiction, template_name)

        if template and isinstance(template, dict):
            body = template.get("body", "")
            if template_name == "rti_request":
                title = "Right to Information Request"
            else:
                title = template.get("title", f"Complaint ({jurisdiction.value})")
        else:
            # Fallback: iterate templates (currently a no-op pass; kept for
            # future expansion of template-lookup logic).
            templates = self.loader.get_templates(jurisdiction)
            for t in templates:
                # templates is a list of dicts from red_flag action references
                pass
            body = "Template not available. Run with LLM client enabled."
            title = "Complaint Draft"

        return ComplaintDraft(
            title=title,
            jurisdiction=jurisdiction,
            body=body,
            template_name=template_name,
            instructions="Fill in the [bracketed] fields with your specific details. "
                         "Review with a qualified attorney before filing.",
        )

    def export_txt(self, draft: ComplaintDraft, path: Optional[str] = None) -> str:
        """Export a complaint draft as a plain-text file.

        Returns the full text content.  If *path* is given, the content is
        also written to that file (parent directories created on demand).
        Default output location is data/output/ when path is None.

        Args:
            draft: The complaint draft to export.
            path: Optional filesystem path to write to.

        Returns:
            The full plain-text content of the draft.
        """
        header = (
            f"{'=' * 60}\n"
            f"  {draft.title}\n"
            f"{'=' * 60}\n\n"
        )
        if draft.template_name:
            header += f"Template: {draft.template_name}\n"
        header += f"Jurisdiction: {draft.jurisdiction.value}\n\n"

        footer = (
            f"\n\n{'-' * 60}\n"
            f"{draft.instructions}\n"
        )

        content = header + draft.body + footer

        if path:
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)

        return content

    def export_docx(self, draft: ComplaintDraft, path: Optional[str] = None) -> str:
        """Export a complaint draft as a .docx file (plain text fallback).

        Uses the ``docx`` library (python-docx) if installed; otherwise falls
        back to saving as a .txt file via :meth:`export_txt`.

        Args:
            draft: The complaint draft to export.
            path: Desired output path.  If None, auto-generates under
                  data/output/.  The extension is normalised to .docx.

        Returns:
            The path to the saved file (either .docx or .txt).
        """
        if path is None:
            path = f"data/output/{draft.jurisdiction.value}_{draft.template_name}.txt"

        try:
            from docx import Document
            doc = Document()
            doc.add_heading(draft.title, level=1)
            doc.add_paragraph(f"Jurisdiction: {draft.jurisdiction.value}")
            if draft.template_name:
                doc.add_paragraph(f"Template: {draft.template_name}")
            doc.add_paragraph("")
            for para in draft.body.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
                else:
                    doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph(draft.instructions)
            outpath = path.replace(".txt", ".docx")
            Path(outpath).parent.mkdir(parents=True, exist_ok=True)
            doc.save(outpath)
            return outpath
        except ImportError:
            # python-docx not installed — save as .txt instead
            txt_path = path.replace(".docx", ".txt")
            self.export_txt(draft, txt_path)
            return txt_path
